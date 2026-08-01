#!/usr/bin/env python3
"""
verify_ats.py — 简历 PDF 的 ATS 文本层与硬性不变量检查

对齐 verify_output.py 的契约化风格：确定性检查（无 LLM），从 config/constraints.yaml
读取 ATS 约束作为单一事实源。用于 Tier2「精投模式」生成 LaTeX CV 后的出厂校验；同时支持 docx 降级路径（python-docx 抽文本层）。

检查项（契约号 [A#]）：
  [A1] 页数 == 期望页数（默认 2）
  [A2] 含可识别联系方式（邮箱 + 电话，字面文本，非仅图标）
  [A3] 文本层无 (cid 乱码 / 替换符（ATS 抽得出字面文本）
  [W-A] 关键词覆盖（JD 原词命中情况，缺失仅警告不硬失败）

文本提取：优先 pypdf（纯 Python，已为项目依赖）；缺失/异常时回退 pdftotext+pdfinfo；docx 用 python-docx 抽段落/表格。

使用方式：
    python3 verify_ats.py --pdf ./cv.pdf [--keywords "Python,分布式,风控"]

退出码：
    0 = 全部通过（含仅 [W] 警告）
    1 = 存在失败断言（[A#]）
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

_SCRIPTS_DIR = Path(__file__).resolve().parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))
from config_loader import load_constraints  # noqa: E402

_CONSTRAINTS = load_constraints()

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(?<!\d)(?:1[3-9]\d{9}|\d{3}-\d{4}-\d{4}|\d{3}-\d{3}-\d{4})(?!\d)")


def read_pdf(pdf_path: str) -> tuple[str, int, str]:
    """返回 (text, page_count, backend)。优先 pypdf，回退 pdftotext+pdfinfo。"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        pages = [p.extract_text() or "" for p in reader.pages]
        return "\n".join(pages), len(reader.pages), "pypdf"
    except Exception:
        pass
    import shutil
    import subprocess
    if shutil.which("pdftotext") and shutil.which("pdfinfo"):
        try:
            txt = subprocess.run(
                ["pdftotext", "-layout", pdf_path, "-"],
                capture_output=True, text=True, timeout=30,
            ).stdout
            info = subprocess.run(
                ["pdfinfo", pdf_path], capture_output=True, text=True, timeout=30,
            ).stdout
            m = re.search(r"Pages:\s*(\d+)", info)
            return txt, int(m.group(1)) if m else 0, "pdftotext"
        except Exception:
            pass
    raise RuntimeError("无法读取 PDF：pypdf 与 pdftotext/pdfinfo 均不可用")


def extract_docx_text(docx_path: str) -> str:
    """用 python-docx 抽取段落与表格文本（供 ATS 文本层检查 / docx 降级路径）。"""
    from docx import Document
    doc = Document(docx_path)
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def read_docx(docx_path: str) -> tuple[str, int, str]:
    """返回 (text, 估算页数, backend)。docx 无原生页码概念，按字符数估算 A4 页。"""
    text = extract_docx_text(docx_path)
    chars_per_page = _CONSTRAINTS.get("ats", {}).get("chars_per_page", 2000)
    page_count = max(1, (len(text) + chars_per_page - 1) // chars_per_page) if text else 1
    return text, page_count, "python-docx"


def run_checks_text(
    text: str,
    jd_keywords: list[str] | None = None,
    page_count: int | None = None,
    constraints: dict | None = None,
    page_check: bool = True,
) -> tuple[list[str], list[str]]:
    """对已经抽取出的文本做 A1-A3/W-A 检查（供 docx / 外部文本层复用）。

    page_count 为 None 时按字符数估算（docx 路径）；page_check=False 时跳过
    [A1] 页数检查（docx 无固定页数语义，仅供 PDF 路径使用）。
    """
    failures: list[str] = []
    warnings: list[str] = []
    cfg = constraints or _CONSTRAINTS
    ats = cfg.get("ats", {})
    chars_per_page = ats.get("chars_per_page", 2000)
    if page_count is None:
        page_count = max(1, (len(text) + chars_per_page - 1) // chars_per_page) if text else 1

    if page_check:
        # [A1] 页数
        expected_pages = ats.get("page_count", 2)
        if page_count != expected_pages:
            failures.append(
                f"[A1] 页数 = {page_count}，期望 {expected_pages}（CV 应严格 {expected_pages} 页；"
                f"文本层提取后端：估算）"
            )

    # [A2] 联系方式（字面文本，非仅图标）
    if ats.get("require_contact", True):
        missing = []
        if not _EMAIL_RE.search(text):
            missing.append("邮箱")
        if not _PHONE_RE.search(text):
            missing.append("电话")
        if missing:
            failures.append(
                f"[A2] 联系方式缺失（{('、'.join(missing))}）——ATS/HR 系统需字面文本，"
                "仅图标 glyph 无效（请确保邮箱/电话以印刷文本出现）"
            )

    # [A3] 乱码
    if ats.get("forbid_garbled", True):
        if "(cid" in text or "\ufffd" in text:
            failures.append(
                "[A3] 文本层含乱码（(cid 或替换符）——PDF 文本层损坏，ATS 抽取会得到垃圾；"
                "检查字体嵌入/编码"
            )

    # [W-A] 关键词覆盖（仅警告，不硬失败）
    if jd_keywords:
        lowered = text.lower()
        covered = [k for k in jd_keywords if k.lower() in lowered]
        missing_kw = [k for k in jd_keywords if k.lower() not in lowered]
        if missing_kw:
            warnings.append(
                f"[W-A] 关键词未命中（用 JD 原词，未添加简历不支持的关键词）：{missing_kw}"
            )
        warnings.append(
            f"[W-A] 关键词覆盖：命中 {len(covered)}/{len(jd_keywords)}；"
            f"缺失 {missing_kw if missing_kw else '无'}"
        )

    return failures, warnings


def run_checks(
    pdf_path: str,
    jd_keywords: list[str] | None = None,
    constraints: dict | None = None,
) -> tuple[list[str], list[str]]:
    """返回 (failures, warnings)。failures 非空 = 退出码 1。

    支持 .pdf（pypdf/pdftotext）与 .docx（python-docx）两种文本层来源。
    客观不变量（联系方式/乱码）为硬失败 [A#]；页数 [A1] 仅对 PDF 校验（docx 无固定页）；
    关键词覆盖为 advisory [W-A]。
    """
    warnings: list[str] = []

    if not Path(pdf_path).exists():
        return [f"[A0] 文件不存在: {pdf_path}"], warnings

    try:
        if str(pdf_path).lower().endswith(".docx"):
            text, page_count, backend = read_docx(pdf_path)
            page_check = False  # docx 无原生页码语义
        else:
            text, page_count, backend = read_pdf(pdf_path)
            page_check = True
    except RuntimeError as e:
        return [f"[A0] {e}"], warnings

    return run_checks_text(
        text, jd_keywords=jd_keywords, page_count=page_count,
        constraints=constraints, page_check=page_check,
    )


def main():
    parser = argparse.ArgumentParser(description="简历 PDF ATS 文本层与硬性不变量检查")
    parser.add_argument("--pdf", required=True, help="简历 PDF 路径")
    parser.add_argument("--keywords", help="JD 关键词，逗号分隔（可选，用于覆盖检查）")
    args = parser.parse_args()

    keywords = [k.strip() for k in args.keywords.split(",") if k.strip()] if args.keywords else None
    failures, warnings = run_checks(args.pdf, keywords)

    if not failures:
        print("✅ ATS 检查通过")
        if warnings:
            print("\n⚠️ 警告（非致命，已显式暴露）：")
            for w in warnings:
                print(f"  {w}")
        sys.exit(0)
    else:
        print(f"❌ {len(failures)} 项 ATS 检查失败:\n")
        for i, f in enumerate(failures, 1):
            print(f"  {i}. {f}")
        if warnings:
            print("\n⚠️ 同时存在的警告：")
            for w in warnings:
                print(f"  {w}")
        sys.exit(1)


if __name__ == "__main__":
    main()
