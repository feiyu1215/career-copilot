#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_cv_docx.py — 非 LaTeX 降级：用 python-docx 从结构化草稿生成 .docx 简历。

当目标机无 LaTeX 引擎（lualatex/xelatex/pdflatex 缺失）时，作为 build_cv.py 的降级路径：

    python drafter_reviewer.py draft --profile p.json --jd jd.txt --format json --out draft.json
    python build_cv.py --draft draft.json --fallback docx --out cv.docx --name ... --email ... --phone ...

本脚本消费 structured JSON（sections + bullets）生成 .docx，并复用 verify_ats 的文本层检查
（联系方式 [A2] / 乱码 [A3] / 关键词覆盖 [W-A]；docx 无固定页数，[A1] 仅对 PDF 校验）。

纯 Python（python-docx），无需 LaTeX / 外部编译；离线可跑。
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

_SCRIPTS_DIR = Path(__file__).resolve().parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

from verify_ats import run_checks_text  # noqa: E402


def _load_draft(draft):
    if isinstance(draft, dict):
        return draft
    p = Path(draft)
    if p.suffix.lower() == ".json":
        return json.loads(p.read_text(encoding="utf-8"))
    return _tex_to_structured(p.read_text(encoding="utf-8"))


def _tex_to_structured(tex: str) -> dict:
    """把 LaTeX 草稿尽力转为单节结构化草稿（降级用，保留可读性文本）。"""
    lines = []
    for raw in tex.splitlines():
        line = raw.strip()
        if not line:
            continue
        line = re.sub(r"\\[a-zA-Z]+\*?\{?", "", line)
        line = re.sub(r"[{}\[\]$\\]", "", line)
        line = line.strip(" %")
        if line:
            lines.append(line)
    return {"sections": [{"title": "正文", "bullets": lines}]}


def _make_docx(sections, *, name="", email="", phone=""):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if name:
        doc.add_heading(name, level=0)
    contact = "  |  ".join([x for x in (email, phone) if x])
    if contact:
        para = doc.add_paragraph()
        run = para.add_run(contact)
        run.font.size = Pt(10)
    for sec in sections or []:
        title = sec.get("title", "经历")
        doc.add_heading(title, level=1)
        bullets = sec.get("bullets") or []
        if bullets:
            for b in bullets:
                doc.add_paragraph(str(b), style="List Bullet")
        else:
            doc.add_paragraph("（无内容）")
    return doc


def build_docx(draft, out_docx=None, *, name="", email="", phone="",
               constraints=None, inspect=True) -> dict:
    """从结构化草稿（dict 或 .json/.tex 路径）生成 .docx。

    返回 {"docx": str, "failures": list, "warnings": list}。
    inspect=True 时跑 verify_ats 文本层检查（docx 路径：跳过 [A1] 页数，保留 [A2]/[A3]/[W-A]）。
    """
    data = _load_draft(draft)
    sections = data.get("sections", []) or []
    out_docx = out_docx or "cv.docx"
    doc = _make_docx(sections, name=name, email=email, phone=phone)
    doc.save(out_docx)

    result = {"docx": str(out_docx), "failures": [], "warnings": []}
    if inspect:
        from verify_ats import extract_docx_text
        text = extract_docx_text(out_docx)
        failures, warnings = run_checks_text(
            text, jd_keywords=None, constraints=constraints, page_check=False)
        result["failures"], result["warnings"] = failures, warnings
    return result


def main():
    ap = argparse.ArgumentParser(description="非 LaTeX 降级：结构化草稿 → .docx 简历")
    ap.add_argument("--draft", required=True, help="draft.json（或 .tex 降级）")
    ap.add_argument("--out", default="cv.docx")
    ap.add_argument("--name", default="")
    ap.add_argument("--email", default="")
    ap.add_argument("--phone", default="")
    ap.add_argument("--no-inspect", action="store_true", help="跳过 ATS 文本层检查")
    args = ap.parse_args()

    res = build_docx(args.draft, args.out, name=args.name, email=args.email,
                     phone=args.phone, inspect=not args.no_inspect)
    print(f"✅ DOCX 简历已生成：{res['docx']}")
    if res["failures"]:
        print(f"❌ ATS 失败 {len(res['failures'])} 项：")
        for f in res["failures"]:
            print(f"  - {f}")
        sys.exit(1)
    if res["warnings"]:
        print("⚠️ ATS 警告：")
        for w in res["warnings"]:
            print(f"  - {w}")
    sys.exit(0)


if __name__ == "__main__":
    main()
